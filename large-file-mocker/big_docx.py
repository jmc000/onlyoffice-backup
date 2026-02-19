#!/usr/bin/env python3
"""
Generate a 12 MB .docx file quickly using images
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from PIL import Image
import io

def create_sample_image(size_kb=100):
    """
    Create a sample PNG image of approximately the specified size in KB.
    
    Args:
        size_kb: Target size in kilobytes
    
    Returns:
        BytesIO object containing the image data
    """
    # Create a colorful gradient image
    # Larger dimensions = larger file size
    width = int((size_kb * 50) ** 0.5)  # Rough estimation
    height = width
    
    img = Image.new('RGB', (width, height))
    pixels = img.load()
    
    # Create a gradient pattern
    for y in range(height):
        for x in range(width):
            r = int((x / width) * 255)
            g = int((y / height) * 255)
            b = int(((x + y) / (width + height)) * 255)
            pixels[x, y] = (r, g, b)
    
    # Save to BytesIO
    img_bytes = io.BytesIO()
    img.save(img_bytes, format='PNG', compress_level=6)
    img_bytes.seek(0)
    
    return img_bytes

def generate_large_docx_fast(output_path, target_size_mb=12, image_path=None):
    """
    Generate a .docx file of approximately the specified size in MB using images.
    
    Args:
        output_path: Path where the .docx file will be saved
        target_size_mb: Target file size in megabytes (default: 12)
        image_path: Optional path to an existing image file (e.g., IMG.png)
    """
    print(f"Generating {target_size_mb} MB .docx file (fast method)...")
    
    # Create a new Document
    doc = Document()
    
    # Add a title with formatting
    title = doc.add_heading('Large Document with Images', 0)
    title_format = title.runs[0]
    title_format.font.color.rgb = RGBColor(0, 51, 102)
    
    # Add description
    intro = doc.add_paragraph(
        f'This document was automatically generated to be approximately {target_size_mb} MB in size. '
        'It uses images to quickly reach the target file size.'
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    target_size_bytes = target_size_mb * 1024 * 1024
    
    # Determine image strategy
    if image_path and os.path.exists(image_path):
        print(f"Using provided image: {image_path}")
        use_provided_image = True
        # Check the image size
        img_size = os.path.getsize(image_path)
        print(f"Image size: {img_size / 1024:.1f} KB")
    else:
        if image_path:
            print(f"Warning: Image '{image_path}' not found. Using generated images instead.")
        print("Generating images programmatically...")
        use_provided_image = False
    
    # Save initial version to check base size
    doc.save(output_path)
    current_size = os.path.getsize(output_path)
    print(f"Initial document size: {current_size / 1024:.1f} KB")
    
    iteration = 0
    
    # Strategy: Add images until we reach target size
    # Images add bulk quickly compared to text
    while current_size < target_size_bytes:
        iteration += 1
        
        # Add a section heading every 5 images
        if iteration % 5 == 0:
            heading = doc.add_heading(f'Image Section {iteration // 5}', level=1)
            heading.runs[0].font.color.rgb = RGBColor(51, 51, 153)
        
        # Add some formatted text
        p = doc.add_paragraph()
        run = p.add_run(f'Image {iteration}: ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        run2 = p.add_run('This is a sample image embedded in the document to increase file size. ')
        run2.font.size = Pt(11)
        run2.italic = True
        
        # Add the image
        try:
            if use_provided_image:
                # Load the image and modify it slightly to make it unique
                # This prevents python-docx from reusing the same image data
                from PIL import Image as PILImage
                img = PILImage.open(image_path)
                
                # Add a small invisible modification (change one pixel)
                # This makes each image unique so it's stored separately
                pixels = img.load()
                if img.mode == 'RGB' or img.mode == 'RGBA':
                    # Modify a single pixel by a tiny amount (imperceptible)
                    x, y = iteration % img.width, (iteration // img.width) % img.height
                    original = pixels[x, y]
                    if img.mode == 'RGB':
                        pixels[x, y] = (original[0], original[1], (original[2] + iteration) % 256)
                    else:  # RGBA
                        pixels[x, y] = (original[0], original[1], (original[2] + iteration) % 256, original[3])
                
                # Save modified image to BytesIO
                img_bytes = io.BytesIO()
                img.save(img_bytes, format='JPEG' if image_path.lower().endswith('.jpeg') or image_path.lower().endswith('.jpg') else 'PNG')
                img_bytes.seek(0)
                
                doc.add_picture(img_bytes, width=Inches(4.0))
            else:
                # Generate a unique sample image each time
                img_data = create_sample_image(size_kb=100)
                doc.add_picture(img_data, width=Inches(3.0))
            
            # Center the image
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception as e:
            print(f"Error adding image: {e}")
            # Add text instead as fallback
            doc.add_paragraph("Image placeholder: " + ("Lorem ipsum " * 100))
        
        # Add a page break every 3 images to prevent overflow
        if iteration % 3 == 0:
            doc.add_page_break()
        
        # Check size periodically
        if iteration % 5 == 0:
            doc.save(output_path)
            current_size = os.path.getsize(output_path)
            progress = (current_size / target_size_bytes) * 100
            print(f"Progress: {progress:.1f}% ({current_size / 1024 / 1024:.2f} MB / {target_size_mb} MB) - {iteration} images added")
            
            # Break if we've exceeded the target
            if current_size >= target_size_bytes:
                break
    
    # Final save
    doc.save(output_path)
    final_size = os.path.getsize(output_path)
    
    print(f"\n✓ Document created successfully!")
    print(f"File: {output_path}")
    print(f"Size: {final_size / 1024 / 1024:.2f} MB ({final_size:,} bytes)")
    print(f"Images added: {iteration}")
    print(f"Pages: ~{iteration // 3}")
    
    return output_path

if __name__ == "__main__":
    import sys
    
    # Check if an image path was provided
    image_file = None
    if len(sys.argv) > 1:
        image_file = sys.argv[1]
        print(f"Using image from command line: {image_file}")
    else:
        # Check for IMG.jpeg or IMG.png in current directory
        if os.path.exists("IMG.jpeg"):
            image_file = "IMG.jpeg"
            print(f"Found IMG.jpeg in current directory")
        elif os.path.exists("IMG.png"):
            image_file = "IMG.png"
            print(f"Found IMG.png in current directory")
    
    # Generate the file in the current directory
    output_file = "large_document_12mb.docx"
    generate_large_docx_fast(output_file, target_size_mb=12, image_path=image_file)
